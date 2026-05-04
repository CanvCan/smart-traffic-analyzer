"""
scripts/generate_docx.py

SYSTEM_DOCUMENTATION.md ve README.md dosyalarini Word (.docx) formatina donusturur.

Kullanim:
    python scripts/generate_docx.py

Cikti:
    Smart_Traffic_Analyzer_Sistem_Dokumantasyonu.docx  (proje kokunde)

Gereksinim:
    pip install python-docx markdown
"""

from pathlib import Path
import re

ROOT = Path(__file__).parent.parent
TRAINING_REPORT    = ROOT / "MODEL_TRAINING_REPORT.md"
INSTALLATION_REPORT = ROOT / "INSTALLATION_REPORT.md"
SYS_DOC            = ROOT / "SYSTEM_DOCUMENTATION.md"
README             = ROOT / "README.md"
OUTPUT             = ROOT / "Smart_Traffic_Analyzer_Sistem_Dokumantasyonu.docx"


def _strip_md(text: str) -> str:
    """Temel Markdown isaretlerini temizler (bold, inline code, link)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text


def build_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("[HATA] python-docx kurulu degil.")
        print("       pip install python-docx")
        return

    doc = Document()

    # ── Sayfa kenarlari ─────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Kapak sayfasi ───────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Smart Traffic Analyzer")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Gerçek Zamanlı Trafik Analiz Sistemi\nSistem Dokümantasyonu")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    doc.add_paragraph()
    doc.add_paragraph()

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        "RF-DETR  ·  ByteTrack  ·  Apache Kafka  ·  Apache Spark  ·  InfluxDB  ·  Grafana"
    )
    info_run.font.size = Pt(11)
    info_run.italic = True

    doc.add_page_break()

    # ── Markdown icerigini isle ─────────────────────────────────────────────
    for md_path in [TRAINING_REPORT, INSTALLATION_REPORT, SYS_DOC, README]:
        if not md_path.exists():
            print(f"[UYARI] Dosya bulunamadi: {md_path}")
            continue

        print(f"[+] Isleniyor: {md_path.name}")
        content = md_path.read_text(encoding="utf-8")
        lines = content.splitlines()

        in_code  = False
        code_buf = []

        for line in lines:
            # ── Kod blogu ────────────────────────────────────────────────
            if line.strip().startswith("```"):
                if in_code:
                    # Kod blogu bitti — yaz
                    code_para = doc.add_paragraph()
                    code_para.style = "No Spacing"
                    code_run = code_para.add_run("\n".join(code_buf))
                    code_run.font.name = "Courier New"
                    code_run.font.size = Pt(8)
                    code_run.font.color.rgb = RGBColor(0x24, 0x24, 0x24)
                    # Arka plan rengini gri yap
                    pPr = code_para._p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    pPr.append(shd)
                    code_buf = []
                    in_code = False
                else:
                    in_code = True
                continue

            if in_code:
                code_buf.append(line)
                continue

            # ── Basliklar ────────────────────────────────────────────────
            if line.startswith("#### "):
                p = doc.add_heading(_strip_md(line[5:].strip()), level=4)
            elif line.startswith("### "):
                p = doc.add_heading(_strip_md(line[4:].strip()), level=3)
            elif line.startswith("## "):
                p = doc.add_heading(_strip_md(line[3:].strip()), level=2)
            elif line.startswith("# "):
                p = doc.add_heading(_strip_md(line[2:].strip()), level=1)

            # ── Tablo satiri ─────────────────────────────────────────────
            elif line.startswith("|"):
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if not cells or all(set(c) <= set("-: ") for c in cells):
                    continue  # ayirici satiri atla
                tbl = doc.add_table(rows=1, cols=len(cells))
                tbl.style = "Table Grid"
                row = tbl.rows[0]
                for i, cell in enumerate(cells):
                    row.cells[i].text = _strip_md(cell)
                    if row.cells[i].paragraphs:
                        run = row.cells[i].paragraphs[0].runs
                        if run:
                            run[0].font.size = Pt(9)

            # ── Listeler ─────────────────────────────────────────────────
            elif line.startswith("- ") or line.startswith("* "):
                text = _strip_md(line[2:].strip())
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(text).font.size = Pt(10)

            elif re.match(r'^\d+\. ', line):
                text = _strip_md(re.sub(r'^\d+\. ', '', line))
                p = doc.add_paragraph(style="List Number")
                p.add_run(text).font.size = Pt(10)

            # ── Yatay cizgi ─────────────────────────────────────────────
            elif line.strip() == "---":
                doc.add_paragraph()

            # ── Normal paragraf ──────────────────────────────────────────
            elif line.strip():
                text = _strip_md(line.strip())
                p = doc.add_paragraph()
                p.add_run(text).font.size = Pt(10)

        # Her dosyadan sonra sayfa sonuna gec
        doc.add_page_break()

    doc.save(str(OUTPUT))
    print(f"\n[OK] Dosya olusturuldu: {OUTPUT}")
    print(f"     Boyut: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_docx()
